"""Resume file upload -> plain text.

Supports PDF (.pdf), Word (.docx), Markdown (.md), plain text (.txt), and Excel
(.xlsx). Decoding is best-effort and fully offline (pypdf / python-docx /
openpyxl); anything that fails degrades to an empty/partial string rather than
crashing the request. The caller (main) passes the raw bytes + the original
filename; this module ignores the extension when decoding is not possible and
falls back to the raw text.

Classes: none. Functions: `decode_resume(data, filename)`.
"""

from __future__ import annotations

import io
import logging
from typing import Any

logger = logging.getLogger("agent.resume_parse")


def _decode_pdf(data: bytes) -> str:
    from pypdf import PdfReader  # lazy import (optional dep)

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - a bad page shouldn't kill the rest
            continue
    return "\n".join(pages)


def _decode_docx(data: bytes) -> str:
    import docx  # lazy import (optional dep)

    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    # include table cells (many resumes encode details as a grid)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _decode_xlsx(data: bytes) -> str:
    import openpyxl  # lazy import (optional dep)

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            if any(v is not None and str(v).strip() for v in row):
                lines.append(" | ".join("" if v is None else str(v).strip() for v in row))
    return "\n".join(lines)


def _decode_md_or_txt(data: bytes) -> str:
    for enc in ("utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="ignore")


def decode_resume(data: bytes, filename: str = "") -> dict[str, Any]:
    """Decode a resume file to plain text.

    Returns {text, format, error}. `format` is one of pdf|docx|md|txt|xlsx|binary.
    Never raises; on failure returns a partial/empty text + an `error` field.
    """
    name = (filename or "").lower()
    text = ""
    fmt = "binary"
    error = ""
    try:
        if name.endswith(".pdf"):
            fmt = "pdf"
            text = _decode_pdf(data)
        elif name.endswith(".docx"):
            fmt = "docx"
            text = _decode_docx(data)
        elif name.endswith(".xlsx") or name.endswith(".xls"):
            fmt = "xlsx"
            text = _decode_xlsx(data)
        elif name.endswith(".md"):
            fmt = "md"
            text = _decode_md_or_txt(data)
        elif name.endswith(".txt"):
            fmt = "txt"
            text = _decode_md_or_txt(data)
        else:
            # Unknown extension: try text; if that's garbage, report binary.
            try:
                text = _decode_md_or_txt(data)
                fmt = "txt"
            except Exception:  # noqa: BLE001
                fmt = "binary"
    except Exception as exc:  # noqa: BLE001
        logger.warning("resume decode failed (%s): %s", name, exc)
        error = str(exc)
    return {"text": (text or "").strip(), "format": fmt, "error": error}
